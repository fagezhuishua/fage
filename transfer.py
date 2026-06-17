#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件类型转换工具 - Modern GUI Edition (customtkinter + tkinterdnd2)
支持：文档/表格/演示/图像/音频/数据格式互转，多语言，拖拽，批量转换
"""

import os
import sys
import shutil
import threading
import json
import csv
import tempfile
import xml.etree.ElementTree as ET
import warnings
import time
import queue
from io import StringIO
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed

import customtkinter as ctk
from customtkinter import (
    CTk, CTkFrame, CTkButton, CTkLabel, CTkTextbox,
    CTkComboBox, CTkSwitch, CTkEntry, CTkProgressBar, CTkToplevel,
)
from tkinter import filedialog, messagebox

# 拖拽支持
try:
    import tkinter
    import tkinterdnd2.TkinterDnD
    from tkinterdnd2 import DND_FILES
    _dnd_attrs = [
        '_subst_format_dnd','_subst_format_str_dnd','_substitute_dnd','_dnd_bind',
        'dnd_bind','drag_source_register','drag_source_unregister',
        'drop_target_register','drop_target_unregister',
        'platform_independent_types','platform_specific_types',
        'get_dropfile_tempdir','set_dropfile_tempdir',
    ]
    for _attr in _dnd_attrs:
        if hasattr(tkinter.BaseWidget, _attr) and not hasattr(tkinter.Tk, _attr):
            setattr(tkinter.Tk, _attr, getattr(tkinter.BaseWidget, _attr))
    _tkdnd_available = True
except ImportError:
    _tkdnd_available = False
    DND_FILES = None

# 可选库
try:
    from docx import Document
except ImportError:
    Document = None
try:
    import pandas as pd
except ImportError:
    pd = None
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except ImportError:
    canvas = None
    letter = None
try:
    from pdf2docx import Converter
except ImportError:
    Converter = None
try:
    from docx2pdf import convert as docx2pdf_convert
except ImportError:
    docx2pdf_convert = None
try:
    import markdown as mdlib
except ImportError:
    mdlib = None
try:
    from PIL import Image, ImageDraw, ImageTk, ImageFont
except ImportError:
    Image = None
    ImageDraw = None
    ImageTk = None
    ImageFont = None
try:
    with warnings.catch_warnings():
        warnings.simplefilter("ignore", RuntimeWarning)
        from pydub import AudioSegment
except ImportError:
    AudioSegment = None
try:
    import win32com.client as win32com_client
except ImportError:
    win32com_client = None

ctk.set_appearance_mode("Light")
ctk.set_default_color_theme("blue")

# ==================== 多语言 ====================
TEXTS = {
    "zh_CN": {
        "app_title": "文件类型转换工具",
        "sidebar_title": "Transfer",
        "sidebar_subtitle": "文件类型转换工具",
        "workbench": "文件类型转换工作台",
        "source_files": "源文件（支持拖拽文件/文件夹到窗口）",
        "add_files": "批量添加文件",
        "add_folder": "添加文件夹",
        "clear_input": "清空输入",
        "clear_output": "清空输出",
        "target_format": "目标格式:",
        "please_add_files": "批量导入入口\n\n1. 点击上方 [批量添加文件] 按钮（可多选）\n2. 点击 [添加文件夹] 按钮（自动递归扫描）\n3. 直接从文件管理器拖拽文件或文件夹到本窗口\n\n支持格式：doc/docx/pdf/txt/xls/xlsx/ppt/pptx/jpg/png 等",
        "output_dir": "输出目录:",
        "browse": "浏览...",
        "convert_result": "转换结果",
        "waiting_convert": "等待转换...",
        "ready": "就绪",
        "format_help": "格式说明",
        "paste_sample": "填入示例",
        "theme": "外观主题",
        "dark_mode": "深色模式",
        "start_convert": "开始转换",
        "hint": "提示",
        "please_add_files_convert": "请先添加待转换的文件",
        "please_select_format": "请选择有效的目标格式",
        "file_not_exist": "文件不存在",
        "error": "错误",
        "saved_to": "已保存到: ",
        "success": "成功",
        "close": "关闭",
        "supported_formats": "支持格式与转换说明",
        "progress": "进度",
        "cleared_input": "已清空输入文件列表",
        "cleared_output": "已清空输出文件",
        "added": "添加",
        "recursive_added": "递归添加",
        "files": "个文件",
        "convert_done": "转换完成，成功",
        "failed_count": "个，失败",
        "open_dir": "打开目录",
        "save_as": "另存为...",
        "skip_unsupported": "跳过不支持的文件: ",
        "language": "语言",
        "restart_required": "语言设置已更改，重启程序后生效。",
        "save": "保存",
        "version": "v3.1 i18n Edition",
        "dragdrop": "拖拽导入",
        "scanning": "正在扫描文件夹...",
    },
    "en": {
        "app_title": "File Transfer",
        "sidebar_title": "Transfer",
        "sidebar_subtitle": "File Type Converter",
        "workbench": "File Conversion Workbench",
        "source_files": "Source Files (Drag and Drop Supported)",
        "add_files": "Batch Add Files",
        "add_folder": "Add Folder",
        "clear_input": "Clear Input",
        "clear_output": "Clear Output",
        "target_format": "Target Format:",
        "please_add_files": "Batch Import\n\n1. Click [Batch Add Files] above (multi-selectable)\n2. Click [Add Folder] (auto recursive scan)\n3. Drag and drop files or folders from explorer\n\nSupported: doc/docx/pdf/txt/xls/xlsx/ppt/pptx/jpg/png etc.",
        "output_dir": "Output Directory:",
        "browse": "Browse...",
        "convert_result": "Conversion Results",
        "waiting_convert": "Waiting for conversion...",
        "ready": "Ready",
        "format_help": "Format Help",
        "paste_sample": "Paste Sample",
        "theme": "Theme",
        "dark_mode": "Dark Mode",
        "start_convert": "Start Conversion",
        "hint": "Hint",
        "please_add_files_convert": "Please add files to convert first",
        "please_select_format": "Please select a valid target format",
        "file_not_exist": "File does not exist",
        "error": "Error",
        "saved_to": "Saved to: ",
        "success": "Success",
        "close": "Close",
        "supported_formats": "Supported Formats and Conversion Info",
        "progress": "Progress",
        "cleared_input": "Input list cleared",
        "cleared_output": "Output files cleared",
        "added": "Added",
        "recursive_added": "Recursively added",
        "files": " files",
        "convert_done": "Conversion completed, success",
        "failed_count": ", failed",
        "open_dir": "Open Dir",
        "save_as": "Save As...",
        "skip_unsupported": "Skipped unsupported file: ",
        "language": "Language",
        "restart_required": "Language changed. Please restart the application.",
        "save": "Save",
        "version": "v3.1 i18n Edition",
        "dragdrop": "Drag and drop imported",
        "scanning": "Scanning folder...",
    },
    "ja": {
        "app_title": "ファイル変換ツール",
        "sidebar_title": "Transfer",
        "sidebar_subtitle": "ファイル変換ツール",
        "workbench": "ファイル変換ワークベンチ",
        "source_files": "ソースファイル（ドラッグ＆ドロップ対応）",
        "add_files": "ファイルを一括追加",
        "add_folder": "フォルダ追加",
        "clear_input": "入力クリア",
        "clear_output": "出力クリア",
        "target_format": "出力形式:",
        "please_add_files": "一括インポート\n\n1. 上部の [ファイルを一括追加] をクリック（複数選択可）\n2. [フォルダ追加] をクリック（自動再帰スキャン）\n3. エクスプローラからファイルやフォルダをドラッグ＆ドロップ\n\n対応形式：doc/docx/pdf/txt/xls/xlsx/ppt/pptx/jpg/png など",
        "output_dir": "出力フォルダ:",
        "browse": "参照...",
        "convert_result": "変換結果",
        "waiting_convert": "変換待ち...",
        "ready": "準備完了",
        "format_help": "形式説明",
        "paste_sample": "サンプル入力",
        "theme": "テーマ",
        "dark_mode": "ダークモード",
        "start_convert": "変換開始",
        "hint": "注意",
        "please_add_files_convert": "変換するファイルを追加してください",
        "please_select_format": "有効な出力形式を選択してください",
        "file_not_exist": "ファイルが存在しません",
        "error": "エラー",
        "saved_to": "保存先: ",
        "success": "成功",
        "close": "閉じる",
        "supported_formats": "対応形式と変換説明",
        "progress": "進捗",
        "cleared_input": "入力リストをクリアしました",
        "cleared_output": "出力ファイルをクリアしました",
        "added": "追加",
        "recursive_added": "再帰追加",
        "files": " ファイル",
        "convert_done": "変換完了、成功",
        "failed_count": "、失敗",
        "open_dir": "フォルダを開く",
        "save_as": "名前を付けて保存...",
        "skip_unsupported": "非対応ファイルをスキップ: ",
        "language": "言語",
        "restart_required": "言語設定が変更されました。アプリを再起動してください。",
        "save": "保存",
        "version": "v3.1 i18n Edition",
        "dragdrop": "ドラッグ＆ドロップで追加",
        "scanning": "フォルダをスキャン中...",
    },
    "de": {
        "app_title": "Dateikonvertierung",
        "sidebar_title": "Transfer",
        "sidebar_subtitle": "Dateikonvertierungstool",
        "workbench": "Konvertierungsarbeitsplatz",
        "source_files": "Quelldateien (Drag and Drop unterstuetzt)",
        "add_files": "Dateien stapeln",
        "add_folder": "Ordner hinzufuegen",
        "clear_input": "Eingabe loeschen",
        "clear_output": "Ausgabe loeschen",
        "target_format": "Zielformat:",
        "please_add_files": "Stapelimport\n\n1. [Dateien stapeln] oben klicken (mehrfachauswahl)\n2. [Ordner hinzufuegen] klicken (automatisch rekursiv)\n3. Dateien oder Ordner aus Explorer hierher ziehen\n\nUnterstuetzt: doc/docx/pdf/txt/xls/xlsx/ppt/pptx/jpg/png usw.",
        "output_dir": "Ausgabeverzeichnis:",
        "browse": "Durchsuchen...",
        "convert_result": "Konvertierungsergebnisse",
        "waiting_convert": "Warte auf Konvertierung...",
        "ready": "Bereit",
        "format_help": "Formatbeschreibung",
        "paste_sample": "Beispiel einfuegen",
        "theme": "Erscheinungsbild",
        "dark_mode": "Dunkelmodus",
        "start_convert": "Konvertierung starten",
        "hint": "Hinweis",
        "please_add_files_convert": "Bitte fuegen Sie zu konvertierende Dateien hinzu",
        "please_select_format": "Bitte waehlen Sie ein gueltiges Zielformat",
        "file_not_exist": "Datei nicht gefunden",
        "error": "Fehler",
        "saved_to": "Gespeichert unter: ",
        "success": "Erfolg",
        "close": "Schliessen",
        "supported_formats": "Unterstuetzte Formate und Konvertierungsinfo",
        "progress": "Fortschritt",
        "cleared_input": "Eingabeliste geloescht",
        "cleared_output": "Ausgabedateien geloescht",
        "added": "Hinzugefuegt",
        "recursive_added": "Rekursiv hinzugefuegt",
        "files": " Dateien",
        "convert_done": "Konvertierung abgeschlossen, Erfolg",
        "failed_count": ", Fehler",
        "open_dir": "Ordner oeffnen",
        "save_as": "Speichern unter...",
        "skip_unsupported": "Nicht unterstuetzte Datei uebersprungen: ",
        "language": "Sprache",
        "restart_required": "Sprache geaendert. Bitte starten Sie die Anwendung neu.",
        "save": "Speichern",
        "version": "v3.1 i18n Edition",
        "dragdrop": "Per Drag and Drop importiert",
        "scanning": "Ordner wird gescannt...",
    },
}

SETTINGS_FILE = "transfer_settings.json"
DEFAULT_SETTINGS = {"language": "zh_CN"}

def _settings_path():
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), SETTINGS_FILE)

def load_settings():
    path = _settings_path()
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                s = json.load(f)
            merged = DEFAULT_SETTINGS.copy()
            merged.update(s)
            return merged
        except Exception:
            pass
    return DEFAULT_SETTINGS.copy()

def save_settings(settings):
    path = _settings_path()
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Save settings failed: {e}")

# ==================== 转换映射 ====================
CONVERSION_MAP = {
    "txt": ["doc","docx","pdf","rtf","odt","html"],
    "doc": ["txt","docx","pdf","rtf","odt","html"],
    "docx": ["txt","doc","pdf","rtf","odt","html","md"],
    "pdf": ["txt","docx","rtf","odt","html","md"],
    "rtf": ["txt","doc","docx","pdf","odt","html"],
    "odt": ["txt","doc","docx","pdf","rtf","html"],
    "md": ["txt","docx","html","pdf","odt"],
    "html": ["txt","doc","docx","pdf","rtf","odt","md"],
    "xls": ["xlsx","csv","html","txt","pdf"],
    "xlsx": ["xls","csv","html","txt","pdf"],
    "csv": ["xls","xlsx","html","txt"],
    "ppt": ["pptx","pdf","html"],
    "pptx": ["ppt","pdf","html"],
    "jpg": ["png","gif","bmp","tiff","webp","svg","ico"],
    "jpeg": ["png","gif","bmp","tiff","webp","svg","ico"],
    "png": ["jpg","gif","bmp","tiff","webp","svg","ico"],
    "gif": ["jpg","png","bmp","tiff","webp","ico"],
    "bmp": ["jpg","png","gif","tiff","webp","ico"],
    "tiff": ["jpg","png","gif","bmp","webp","ico"],
    "webp": ["jpg","png","gif","bmp","tiff","ico"],
    "svg": ["png","jpg","bmp"],
    "wav": ["mp3","ogg","flac","aiff","au"],
    "mp3": ["wav","ogg","flac","aiff","au"],
    "ogg": ["wav","mp3","flac","aiff","au"],
    "flac": ["wav","mp3","ogg","aiff","au"],
    "aiff": ["wav","mp3","ogg","flac","au"],
    "au": ["wav","mp3","ogg","flac","aiff"],
    "xml": ["json","csv","txt","html"],
    "json": ["xml","csv","txt","html"],
}

ALL_TARGET_FORMATS = sorted({fmt for targets in CONVERSION_MAP.values() for fmt in targets})

FORMAT_DESCRIPTIONS = {
    "txt": "纯文本文档，支持所有文档转换",
    "doc": "Microsoft Word 97-2003 文档，需安装 Office",
    "docx": "Microsoft Word 2007+ 文档，推荐格式",
    "pdf": "可移植文档格式，适合阅读和打印",
    "rtf": "富文本文档，跨平台兼容性好",
    "odt": "开放文档格式，开源办公套件支持",
    "md": "Markdown 标记语言，适合技术文档",
    "html": "网页文档，可在浏览器中查看",
    "xls": "Excel 97-2003 电子表格，需安装 Office",
    "xlsx": "Excel 2007+ 电子表格，推荐格式",
    "csv": "逗号分隔值，数据交换常用格式",
    "ppt": "PowerPoint 97-2003 演示文稿，需安装 Office",
    "pptx": "PowerPoint 2007+ 演示文稿，推荐格式",
    "jpg": "JPEG 图像，有损压缩，适合照片",
    "jpeg": "同 jpg",
    "png": "PNG 图像，无损压缩，支持透明",
    "gif": "GIF 图像，支持动画和透明",
    "bmp": "BMP 位图，无压缩，文件较大",
    "tiff": "TIFF 图像，高质量，适合印刷",
    "webp": "WebP 图像，现代高效格式",
    "svg": "SVG 矢量图，无损缩放",
    "wav": "WAV 音频，无压缩高质量",
    "mp3": "MP3 音频，有损压缩，通用格式",
    "ogg": "OGG 音频，开源免费",
    "flac": "FLAC 音频，无损压缩",
    "aiff": "AIFF 音频，专业音频格式",
    "au": "AU 音频，Unix/Linux 常用",
    "xml": "可扩展标记语言，结构化数据",
    "json": "JavaScript 对象表示法，轻量数据格式",
    "ico": "Windows 图标文件",
}


class ConversionError(Exception):
    UNSUPPORTED = "UNSUPPORTED"
    MISSING_DEP = "MISSING_DEP"
    FILE_NOT_FOUND = "FILE_NOT_FOUND"
    EMPTY = "EMPTY"
    CORRUPTED = "CORRUPTED"
    PERMISSION = "PERMISSION"
    FAILED = "FAILED"

    def __init__(self, error_type, message="", details=""):
        self.error_type = error_type
        self.message = message
        self.details = details
        super().__init__(self.message)

    def __str__(self):
        return f"[{self.error_type}] {self.message}: {self.details}"


def get_allowed_formats(source_ext):
    ext = source_ext.lower().lstrip(".")
    return CONVERSION_MAP.get(ext, [])


def _get_ext(filepath):
    return os.path.splitext(filepath)[1].lower().lstrip(".")


def convert_with_com(filepath, out_dir, target_format):
    if not win32com_client:
        raise ConversionError(ConversionError.MISSING_DEP, "缺少 win32com.client", "pip install pywin32")
    abs_path = os.path.abspath(filepath)
    if not os.path.exists(abs_path):
        raise ConversionError(ConversionError.FILE_NOT_FOUND, filepath)
    base_name = os.path.splitext(os.path.basename(filepath))[0]
    out_path = os.path.join(out_dir, f"{base_name}.{target_format}")
    ext = _get_ext(filepath)
    try:
        if ext in ("doc","docx","rtf","odt"):
            if target_format not in ("txt","pdf","doc","docx","rtf","html","odt"):
                raise ConversionError(ConversionError.UNSUPPORTED, f"不支持 {ext} -> {target_format}")
            word = win32com_client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(abs_path)
            fmt_map = {"txt":2,"pdf":17,"doc":0,"docx":16,"rtf":6,"html":8,"odt":23}
            doc.SaveAs2(out_path, fmt_map.get(target_format, 16))
            doc.Close()
            word.Quit()
        elif ext in ("xls","xlsx"):
            if target_format not in ("csv","html","txt","xls","xlsx","pdf"):
                raise ConversionError(ConversionError.UNSUPPORTED, f"不支持 {ext} -> {target_format}")
            excel = win32com_client.Dispatch("Excel.Application")
            excel.Visible = False
            wb = excel.Workbooks.Open(abs_path)
            if target_format == "pdf":
                wb.ExportAsFixedFormat(0, out_path)
            else:
                fmt_map = {"csv":6,"html":44,"txt":-4158,"xls":56,"xlsx":51}
                wb.SaveAs(out_path, fmt_map.get(target_format, 51))
            wb.Close()
            excel.Quit()
        elif ext in ("ppt","pptx"):
            if target_format not in ("pdf","html","ppt","pptx"):
                raise ConversionError(ConversionError.UNSUPPORTED, f"不支持 {ext} -> {target_format}")
            ppt = win32com_client.Dispatch("PowerPoint.Application")
            ppt.Visible = False
            prs = ppt.Presentations.Open(abs_path)
            if target_format == "pdf":
                prs.SaveAs(out_path, 32)
            elif target_format == "html":
                prs.SaveAs(out_path, 12)
            else:
                fmt_map = {"ppt":1,"pptx":11}
                prs.SaveAs(out_path, fmt_map.get(target_format, 11))
            prs.Close()
            ppt.Quit()
        else:
            raise ConversionError(ConversionError.UNSUPPORTED, f"不支持 COM 转换: {ext}")
        return out_path
    except ConversionError:
        raise
    except Exception as e:
        raise ConversionError(ConversionError.FAILED, str(e), f"COM 转换失败: {ext} -> {target_format}")


def _text_to_docx(text_content, out_path):
    """将纯文本内容保存为 docx"""
    if Document is None:
        raise ConversionError(ConversionError.MISSING_DEP, "缺少 python-docx")
    doc = Document()
    for line in text_content.splitlines():
        doc.add_paragraph(line)
    doc.save(out_path)


def convert_file(filepath, out_dir, target_format, status_callback=None):
    ext = _get_ext(filepath)
    target_format = target_format.lower().lstrip(".")

    if ext == target_format:
        out_path = os.path.join(out_dir, os.path.basename(filepath))
        shutil.copy2(filepath, out_path)
        return out_path

    allowed = get_allowed_formats(ext)
    if target_format not in allowed:
        raise ConversionError(
            ConversionError.UNSUPPORTED,
            f"不支持 {ext} 到 {target_format} 的转换",
            f"支持的格式: {', '.join(allowed)}"
        )

    base_name = os.path.splitext(os.path.basename(filepath))[0]
    out_path = os.path.join(out_dir, f"{base_name}.{target_format}")
    os.makedirs(out_dir, exist_ok=True)

    if status_callback:
        status_callback(f"正在转换 {os.path.basename(filepath)} -> {target_format}...")

    try:
        # ========== 纯文本/标记语言 ↔ 文档 ==========
        if ext in ("txt","md","html","xml","json"):
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f_in:
                content = f_in.read()
            # 纯文本互转
            if target_format in ("txt","html","md"):
                with open(out_path, "w", encoding="utf-8") as f_out:
                    f_out.write(content)
                return out_path
            # → docx
            elif target_format == "docx":
                _text_to_docx(content, out_path)
                return out_path
            # → doc / rtf / odt（先 docx 再 COM 转）
            elif target_format in ("doc","rtf","odt"):
                if not win32com_client:
                    raise ConversionError(ConversionError.MISSING_DEP, f"{ext} → {target_format} 需要安装 Office")
                tmp_docx = os.path.join(tempfile.gettempdir(), f"_{base_name}_{os.getpid()}.docx")
                _text_to_docx(content, tmp_docx)
                try:
                    return convert_with_com(tmp_docx, out_dir, target_format)
                finally:
                    if os.path.exists(tmp_docx):
                        os.remove(tmp_docx)
            # → pdf（先 docx 再转 pdf）
            elif target_format == "pdf":
                tmp_docx = os.path.join(tempfile.gettempdir(), f"_{base_name}_{os.getpid()}.docx")
                _text_to_docx(content, tmp_docx)
                try:
                    if docx2pdf_convert:
                        docx2pdf_convert(tmp_docx, out_path)
                        return out_path
                    if win32com_client:
                        return convert_with_com(tmp_docx, out_dir, "pdf")
                    raise ConversionError(ConversionError.MISSING_DEP, "需要 docx2pdf 或 Office")
                finally:
                    if os.path.exists(tmp_docx):
                        os.remove(tmp_docx)

        # ========== docx ↔ 其他 ==========
        elif ext == "docx":
            if target_format == "txt":
                if not Document:
                    raise ConversionError(ConversionError.MISSING_DEP, "缺少 python-docx")
                doc = Document(filepath)
                text = "\n".join([p.text for p in doc.paragraphs])
                with open(out_path, "w", encoding="utf-8") as f_out:
                    f_out.write(text)
                return out_path
            elif target_format == "pdf":
                if docx2pdf_convert:
                    docx2pdf_convert(filepath, out_path)
                    return out_path
                if win32com_client:
                    return convert_with_com(filepath, out_dir, "pdf")
                raise ConversionError(ConversionError.MISSING_DEP, "缺少 docx2pdf 或 Office")
            elif target_format in ("doc","html","rtf","odt"):
                if not win32com_client:
                    raise ConversionError(ConversionError.MISSING_DEP, "需要安装 Office")
                return convert_with_com(filepath, out_dir, target_format)

        # ========== doc ↔ 其他 ==========
        elif ext == "doc":
            if target_format in ("docx","txt","html","rtf","odt","pdf"):
                if not win32com_client:
                    raise ConversionError(ConversionError.MISSING_DEP, "需要安装 Office")
                return convert_with_com(filepath, out_dir, target_format)

        # ========== rtf / odt ↔ 其他 ==========
        elif ext in ("rtf","odt"):
            if target_format in ("txt","doc","docx","pdf","html","rtf","odt"):
                if not win32com_client:
                    raise ConversionError(ConversionError.MISSING_DEP, "需要安装 Office")
                return convert_with_com(filepath, out_dir, target_format)
            raise ConversionError(ConversionError.UNSUPPORTED, f"不支持 {ext} -> {target_format}")

        # ========== pdf ↔ 其他 ==========
        elif ext == "pdf":
            if target_format == "docx":
                if Converter:
                    cv = Converter(filepath)
                    cv.convert(out_path, start=0, end=None)
                    cv.close()
                    return out_path
                raise ConversionError(ConversionError.MISSING_DEP, "缺少 pdf2docx")
            elif target_format == "txt":
                if PdfReader:
                    reader = PdfReader(filepath)
                    text = "\n".join([page.extract_text() or "" for page in reader.pages])
                    with open(out_path, "w", encoding="utf-8") as f_out:
                        f_out.write(text)
                    return out_path
                raise ConversionError(ConversionError.MISSING_DEP, "缺少 PyPDF2")
            elif target_format in ("doc","rtf","odt","html","md"):
                if not Converter:
                    raise ConversionError(ConversionError.MISSING_DEP, "缺少 pdf2docx")
                if not win32com_client:
                    raise ConversionError(ConversionError.MISSING_DEP, "需要安装 Office")
                tmp_docx = os.path.join(tempfile.gettempdir(), f"_{base_name}_{os.getpid()}.docx")
                cv = Converter(filepath)
                cv.convert(tmp_docx, start=0, end=None)
                cv.close()
                try:
                    return convert_with_com(tmp_docx, out_dir, target_format)
                finally:
                    if os.path.exists(tmp_docx):
                        os.remove(tmp_docx)

        # 表格处理
        elif ext in ("xls","xlsx","csv"):
            if pd is None:
                raise ConversionError(ConversionError.MISSING_DEP, "缺少 pandas")
            try:
                if ext == "csv":
                    df = pd.read_csv(filepath)
                else:
                    df = pd.read_excel(filepath)
            except Exception as e:
                raise ConversionError(ConversionError.FAILED, str(e), f"读取 {ext} 失败，可能需要安装 openpyxl/xlrd")
            if target_format == "csv":
                df.to_csv(out_path, index=False, encoding="utf-8-sig")
                return out_path
            elif target_format in ("xls","xlsx"):
                try:
                    if target_format == "xls":
                        df.to_excel(out_path, index=False, engine="xlwt")
                    else:
                        df.to_excel(out_path, index=False, engine="openpyxl")
                    return out_path
                except Exception as e:
                    raise ConversionError(ConversionError.FAILED, str(e), f"保存 {target_format} 失败，请安装 openpyxl/xlwt")
            elif target_format == "html":
                df.to_html(out_path, index=False)
                return out_path
            elif target_format == "txt":
                with open(out_path, "w", encoding="utf-8") as f_out:
                    f_out.write(df.to_string(index=False))
                return out_path
            elif target_format == "pdf":
                if win32com_client:
                    return convert_with_com(filepath, out_dir, target_format)
                raise ConversionError(ConversionError.MISSING_DEP, "需要安装 Office")
            else:
                raise ConversionError(ConversionError.UNSUPPORTED, f"表格不支持转换为 {target_format}")

        # 演示处理
        elif ext in ("ppt","pptx"):
            if win32com_client:
                return convert_with_com(filepath, out_dir, target_format)
            raise ConversionError(ConversionError.MISSING_DEP, "需要安装 Office")

        # 图像处理
        elif ext in ("jpg","jpeg","png","gif","bmp","tiff","webp","svg"):
            if Image is None:
                raise ConversionError(ConversionError.MISSING_DEP, "缺少 Pillow")
            # SVG 需要特殊处理
            if ext == "svg":
                try:
                    import cairosvg
                    png_tmp = os.path.join(tempfile.gettempdir(), f"_{base_name}_{os.getpid()}.png")
                    cairosvg.svg2png(url=filepath, write_to=png_tmp)
                    if target_format == "png":
                        shutil.copy2(png_tmp, out_path)
                        os.remove(png_tmp)
                        return out_path
                    img = Image.open(png_tmp)
                    img.save(out_path)
                    os.remove(png_tmp)
                    return out_path
                except ImportError:
                    raise ConversionError(ConversionError.MISSING_DEP, "SVG 转换需要安装 cairosvg")
                except Exception as e:
                    raise ConversionError(ConversionError.FAILED, str(e), "SVG 转换失败")
            img = Image.open(filepath)
            # ICO 格式需要特殊处理
            if target_format == "ico":
                if img.mode in ("RGBA","P"):
                    img = img.convert("RGB")
                elif img.mode == "LA":
                    img = img.convert("L")
                # 缩放为常用尺寸并保存
                sizes = [(16,16),(32,32),(48,48),(64,64),(128,128),(256,256)]
                icons = []
                for size in sizes:
                    icons.append(img.resize(size, Image.Resampling.LANCZOS))
                icons[0].save(out_path, format="ICO", sizes=[i.size for i in icons])
                return out_path
            # 其他格式
            if img.mode in ("RGBA","P"):
                img = img.convert("RGB")
            elif img.mode == "LA":
                img = img.convert("L")
            img.save(out_path)
            return out_path

        # 音频处理
        elif ext in ("wav","mp3","ogg","flac","aiff","au"):
            if AudioSegment is None:
                raise ConversionError(ConversionError.MISSING_DEP, "缺少 pydub（ffmpeg 可选）")
            audio = AudioSegment.from_file(filepath, format=ext)
            audio.export(out_path, format=target_format)
            return out_path

        # 数据格式处理
        elif ext in ("xml","json"):
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f_in:
                content = f_in.read()
            if target_format == "csv":
                try:
                    if ext == "json":
                        data = json.loads(content)
                        if isinstance(data, list) and data and isinstance(data[0], dict):
                            keys = list(data[0].keys())
                            with open(out_path, "w", newline="", encoding="utf-8-sig") as f_out:
                                writer = csv.DictWriter(f_out, fieldnames=keys)
                                writer.writeheader()
                                writer.writerows(data)
                            return out_path
                        else:
                            raise ValueError("JSON 不是对象列表，无法转为 CSV")
                    elif ext == "xml":
                        root = ET.fromstring(content)
                        rows = []
                        for child in root:
                            row = {elem.tag: elem.text or "" for elem in child}
                            rows.append(row)
                        if rows:
                            keys = list(rows[0].keys())
                            with open(out_path, "w", newline="", encoding="utf-8-sig") as f_out:
                                writer = csv.DictWriter(f_out, fieldnames=keys)
                                writer.writeheader()
                                writer.writerows(rows)
                            return out_path
                        raise ValueError("XML 中没有可解析的数据")
                except Exception as e:
                    raise ConversionError(ConversionError.FAILED, str(e), f"{ext} -> CSV 失败")
            elif target_format in ("txt","html"):
                with open(out_path, "w", encoding="utf-8") as f_out:
                    f_out.write(content)
                return out_path
            elif target_format == "xml" and ext == "json":
                try:
                    data = json.loads(content)
                    root = ET.Element("root")
                    _dict_to_xml(data, root)
                    tree = ET.ElementTree(root)
                    tree.write(out_path, encoding="utf-8", xml_declaration=True)
                    return out_path
                except Exception as e:
                    raise ConversionError(ConversionError.FAILED, str(e), "JSON -> XML 失败")
            elif target_format == "json" and ext == "xml":
                try:
                    root = ET.fromstring(content)
                    data = _xml_to_dict(root)
                    with open(out_path, "w", encoding="utf-8") as f_out:
                        json.dump(data, f_out, ensure_ascii=False, indent=2)
                    return out_path
                except Exception as e:
                    raise ConversionError(ConversionError.FAILED, str(e), "XML -> JSON 失败")
            else:
                raise ConversionError(ConversionError.UNSUPPORTED, f"不支持 {ext} -> {target_format}")

        else:
            raise ConversionError(ConversionError.UNSUPPORTED, f"不支持 {ext} -> {target_format}")

    except ConversionError:
        raise
    except Exception as e:
        import traceback
        raise ConversionError(ConversionError.FAILED, str(e), traceback.format_exc())


def _dict_to_xml(data, parent):
    if isinstance(data, dict):
        for key, val in data.items():
            child = ET.SubElement(parent, str(key))
            _dict_to_xml(val, child)
    elif isinstance(data, list):
        for item in data:
            child = ET.SubElement(parent, "item")
            _dict_to_xml(item, child)
    else:
        parent.text = str(data) if data is not None else ""


def _xml_to_dict(element):
    result = {}
    for child in element:
        if len(child) == 0:
            result[child.tag] = child.text or ""
        else:
            result[child.tag] = _xml_to_dict(child)
    return result


class FileTransferApp(CTk):
    def __init__(self):
        super().__init__()
        self.settings = load_settings()
        self.lang = self.settings.get("language", "zh_CN")
        self._configure_window()
        self._setup_background()
        self._init_ui()
        self._setup_drag_drop()

    def _t(self, key):
        return TEXTS.get(self.lang, TEXTS["zh_CN"]).get(key, key)

    def _configure_window(self):
        self.title(self._t("app_title"))
        self.geometry("1200x800")
        self.minsize(1000, 700)

    def _generate_background(self, width=1200, height=800):
        if Image is None or ImageDraw is None:
            return None
        cache_key = (width, height, ctk.get_appearance_mode())
        if hasattr(self, "_bg_cache") and self._bg_cache.get("key") == cache_key:
            return self._bg_cache["image"]
        is_dark = ctk.get_appearance_mode() == "Dark"
        # 使用 PIL 的渐变方式比逐像素 line 更快
        img = Image.new("RGB", (width, height))
        draw = ImageDraw.Draw(img)
        if is_dark:
            top_color = (18, 22, 50)
            bot_color = (28, 30, 70)
            deco_color = (45, 55, 100)
            text_color = (35, 42, 75)
        else:
            top_color = (232, 240, 248)
            bot_color = (245, 248, 255)
            deco_color = (190, 200, 220)
            text_color = (170, 180, 205)
        # 使用 numpy 加速渐变（如果可用），否则用 rectangle 分块
        try:
            import numpy as np
            arr = np.zeros((height, width, 3), dtype=np.uint8)
            for c in range(3):
                col_start = top_color[c]
                col_end = bot_color[c]
                arr[:, :, c] = np.linspace(col_start, col_end, height, dtype=np.uint8)[:, np.newaxis]
            img = Image.fromarray(arr)
            draw = ImageDraw.Draw(img)
        except Exception:
            # 回退：用渐变矩形条，每 4 像素一条，减少 draw 调用
            step = 4
            for y in range(0, height, step):
                ratio = y / height
                r = int(top_color[0] + (bot_color[0] - top_color[0]) * ratio)
                g = int(top_color[1] + (bot_color[1] - top_color[1]) * ratio)
                b = int(top_color[2] + (bot_color[2] - top_color[2]) * ratio)
                draw.rectangle([0, y, width, min(y + step, height)], fill=(r, g, b))
        import random
        random.seed(42)
        formats = ["PDF", "DOCX", "XLSX", "JPG", "PNG", "MP3", "TXT", "CSV", "PPTX", "XML", "JSON", "HTML", "RTF", "ODT", "WAV"]
        for _ in range(20):
            x = random.randint(0, width)
            y = random.randint(0, height)
            fmt = random.choice(formats)
            draw.text((x, y), fmt, fill=text_color)
        for _ in range(10):
            x = random.randint(0, width - 40)
            y = random.randint(0, height - 50)
            draw.rectangle([x, y, x + 30, y + 40], outline=deco_color, width=1)
            draw.polygon([(x + 20, y), (x + 30, y), (x + 30, y + 10)], fill=deco_color)
        for _ in range(8):
            x = random.randint(0, width - 30)
            y = random.randint(0, height - 10)
            draw.line([(x, y), (x + 20, y)], fill=deco_color, width=2)
            draw.polygon([(x + 20, y - 5), (x + 30, y), (x + 20, y + 5)], fill=deco_color)
        self._bg_cache = {"key": cache_key, "image": img}
        return img

    def _setup_background(self):
        if ImageTk is None:
            return
        self.update_idletasks()
        width = max(self.winfo_width(), 1200)
        height = max(self.winfo_height(), 800)
        bg_pil = self._generate_background(width, height)
        if bg_pil is None:
            return
        self._bg_image = ImageTk.PhotoImage(bg_pil)
        if hasattr(self, "_bg_label"):
            self._bg_label.destroy()
        self._bg_label = tkinter.Label(self, image=self._bg_image)
        self._bg_label.place(x=0, y=0, relwidth=1, relheight=1)
        self._bg_label.lower()

    def _init_ui(self):
        # 侧边栏
        self.sidebar = CTkFrame(self, width=200, corner_radius=0, fg_color=("#e8edf5", "#0f142d"))
        self.sidebar.pack(side="left", fill="y", padx=(0, 8))
        self.sidebar.pack_propagate(False)

        CTkLabel(self.sidebar, text=self._t("sidebar_title"), font=("Microsoft YaHei", 24, "bold"), text_color=("#2c3e50", "#e0e6f0")).pack(pady=30)
        CTkLabel(self.sidebar, text=self._t("sidebar_subtitle"), font=("Microsoft YaHei", 12), text_color=("#5a6a7a", "#8899aa")).pack()

        CTkLabel(self.sidebar, text=self._t("theme"), font=("Microsoft YaHei", 12, "bold")).pack(pady=(20, 5))
        self.theme_switch = CTkSwitch(self.sidebar, text=self._t("dark_mode"), command=self._toggle_theme)
        self.theme_switch.pack(pady=5)
        # 默认 Light 模式，switch 不选中

        CTkLabel(self.sidebar, text=self._t("language"), font=("Microsoft YaHei", 12, "bold")).pack(pady=(20, 5))
        self.lang_combo = CTkComboBox(self.sidebar, values=["zh_CN", "en", "ja", "de"], command=self._change_language)
        self.lang_combo.set(self.lang)
        self.lang_combo.pack(pady=5)

        CTkButton(self.sidebar, text=self._t("format_help"), command=self._show_format_help).pack(pady=10, padx=20, fill="x")
        CTkButton(self.sidebar, text=self._t("paste_sample"), command=self._paste_sample).pack(pady=5, padx=20, fill="x")

        CTkLabel(self.sidebar, text=self._t("version"), font=("Microsoft YaHei", 10), text_color="gray").pack(side="bottom", pady=20)

        # 主区域
        self.main_frame = CTkFrame(self, fg_color=("#eef2f7", "#12183a"))
        self.main_frame.pack(side="left", fill="both", expand=True, padx=(0, 20), pady=20)

        CTkLabel(self.main_frame, text=self._t("workbench"), font=("Microsoft YaHei", 18, "bold"), text_color=("#1a2530", "#d0d8e8")).pack(anchor="w", pady=(0, 15))

        # 源文件区域
        src_frame = CTkFrame(self.main_frame, fg_color=("#f5f7fa", "#1a1f4b"))
        src_frame.pack(fill="both", expand=True, pady=(0, 10))

        src_header = CTkFrame(src_frame, fg_color="transparent")
        src_header.pack(fill="x", padx=10, pady=(10, 5))

        CTkLabel(src_header, text=self._t("source_files"), font=("Microsoft YaHei", 14, "bold"), text_color=("#2c3e50", "#c8d4e8")).pack(side="left")

        CTkButton(src_header, text=self._t("add_files"), command=self._add_files, width=100).pack(side="right", padx=5)
        CTkButton(src_header, text=self._t("add_folder"), command=self._add_folder, width=100).pack(side="right", padx=5)
        CTkButton(src_header, text=self._t("clear_input"), command=self._clear_input, width=120).pack(side="right", padx=5)

        self.files_list = CTkTextbox(src_frame, wrap="word", height=150, fg_color=("#ffffff", "#232850"), text_color=("#1a1a1a", "#e0e0e0"))
        self.files_list.pack(fill="both", expand=True, padx=10, pady=(0, 10))
        self.files_list.insert("1.0", self._t("please_add_files"))
        self.files_list.configure(state="disabled")

        self.source_files = []
        self.converted_files = []

        # 配置行
        cfg_row = CTkFrame(self.main_frame, fg_color=("#f5f7fa", "#1a1f4b"))
        cfg_row.pack(fill="x", pady=5)

        CTkLabel(cfg_row, text=self._t("target_format"), font=("Microsoft YaHei", 12), text_color=("#2c3e50", "#c8d4e8")).pack(side="left", padx=5)
        self.target_format = CTkComboBox(cfg_row, values=ALL_TARGET_FORMATS, width=120)
        self.target_format.pack(side="left", padx=5)
        self.target_format.set("")

        self.start_btn = CTkButton(cfg_row, text=self._t("start_convert"), command=self._start_convert, fg_color="#2E7D32", hover_color="#1B5E20", width=140)
        self.start_btn.pack(side="left", padx=10)

        self.clear_output_btn = CTkButton(cfg_row, text=self._t("clear_output"), command=self._clear_output, width=120)
        self.clear_output_btn.pack(side="right", padx=5)

        CTkLabel(cfg_row, text=self._t("output_dir"), font=("Microsoft YaHei", 12), text_color=("#2c3e50", "#c8d4e8")).pack(side="right", padx=5)
        self.output_entry = CTkEntry(cfg_row, placeholder_text=self._get_default_output_dir())
        self.output_entry.pack(side="right", fill="x", expand=True, padx=5)
        self.output_entry.insert(0, self._get_default_output_dir())

        CTkButton(cfg_row, text=self._t("browse"), command=self._browse_output, width=80).pack(side="right", padx=5)

        # 进度条
        self.progress = CTkProgressBar(self.main_frame)
        self.progress.pack(fill="x", pady=10)
        self.progress.set(0)

        self.progress_label = CTkLabel(self.main_frame, text=self._t("ready"), text_color=("#4a5568", "#a0aec0"))
        self.progress_label.pack(anchor="w")

        # 转换结果
        CTkLabel(self.main_frame, text=self._t("convert_result"), font=("Microsoft YaHei", 14, "bold"), text_color=("#2c3e50", "#c8d4e8")).pack(anchor="w", pady=(10, 5))

        result_frame = CTkFrame(self.main_frame, fg_color=("#f5f7fa", "#1a1f4b"))
        result_frame.pack(fill="both", expand=True)

        self.result_text = CTkTextbox(result_frame, wrap="word", fg_color=("#ffffff", "#232850"), text_color=("#1a1a1a", "#e0e0e0"))
        self.result_text.pack(fill="both", expand=True, padx=10, pady=10)
        self.result_text.insert("1.0", self._t("waiting_convert"))
        self.result_text.configure(state="disabled")

    def _get_default_output_dir(self):
        return os.path.join(os.path.expanduser("~"), "Downloads", "transfer_output")

    def _toggle_theme(self):
        mode = "Dark" if self.theme_switch.get() else "Light"
        ctk.set_appearance_mode(mode)
        self.after(100, self._setup_background)

    def _change_language(self, choice):
        self.settings["language"] = choice
        save_settings(self.settings)
        messagebox.showinfo(self._t("hint"), self._t("restart_required"))

    def _setup_drag_drop(self):
        if not _tkdnd_available:
            return
        try:
            tkinterdnd2.TkinterDnD._require(self)
            self.drop_target_register(DND_FILES)
            self.dnd_bind("<<Drop>>", self._on_drop)
        except Exception as e:
            print(f"DnD setup failed: {e}")

    def _on_drop(self, event):
        data = event.data
        if not data:
            return
        files = self._parse_drop(data)
        # 立即添加文件，目录异步扫描
        immediate_files = [fp for fp in files if os.path.isfile(fp)]
        dirs = [fp for fp in files if os.path.isdir(fp)]
        self.source_files.extend(immediate_files)
        if dirs:
            self.progress_label.configure(text=self._t("scanning"))
            def _scan_dirs():
                new_files = []
                for d in dirs:
                    new_files.extend(self._scan_dir(d))
                self.after(0, lambda: self._on_drop_scanned(new_files))
            threading.Thread(target=_scan_dirs, daemon=True).start()
        if immediate_files or not dirs:
            self._update_files_list()

    def _on_drop_scanned(self, new_files):
        self.source_files.extend(new_files)
        self._update_files_list()
        self.progress_label.configure(text=self._t("ready"))

    def _parse_drop(self, data):
        if not data:
            return []
        parts = []
        current = ""
        i = 0
        while i < len(data):
            if data[i] == '{':
                j = data.find('}', i)
                if j != -1:
                    parts.append(data[i+1:j])
                    i = j + 1
                else:
                    current += data[i]
                    i += 1
            elif data[i] == ' ':
                if current:
                    parts.append(current)
                    current = ""
                i += 1
            else:
                current += data[i]
                i += 1
        if current:
            parts.append(current)
        return [p.strip().strip('"').strip("'") for p in parts if p.strip()]

    def _scan_dir(self, directory, recursive=True):
        result = []
        for root, dirs, files in os.walk(directory):
            for f in files:
                fp = os.path.join(root, f)
                if os.path.isfile(fp):
                    result.append(fp)
            if not recursive:
                break
        return result

    def _update_files_list(self):
        self.files_list.configure(state="normal")
        self.files_list.delete("1.0", "end")
        if self.source_files:
            # 批量插入，减少刷新开销
            chunk_size = 500
            total = len(self.source_files)
            for i in range(0, total, chunk_size):
                chunk = self.source_files[i:i+chunk_size]
                self.files_list.insert("end", "\n".join(chunk) + "\n")
        else:
            self.files_list.insert("1.0", self._t("please_add_files"))
        self.files_list.configure(state="disabled")

    def _add_files(self):
        files = filedialog.askopenfilenames(title=self._t("add_files"))
        if files:
            self.source_files.extend(files)
            self._update_files_list()

    def _add_folder(self):
        folder = filedialog.askdirectory(title=self._t("add_folder"))
        if folder:
            self.progress_label.configure(text=self._t("scanning"))
            self.after(10, lambda: self._add_folder_async(folder))

    def _add_folder_async(self, folder):
        def _scan():
            files = self._scan_dir(folder)
            self.after(0, lambda: self._on_folder_scanned(files))
        threading.Thread(target=_scan, daemon=True).start()

    def _on_folder_scanned(self, files):
        self.source_files.extend(files)
        self._update_files_list()
        self.progress_label.configure(text=self._t("ready"))

    def _clear_input(self):
        self.source_files = []
        self._update_files_list()
        self.progress.set(0)
        self.progress_label.configure(text=self._t("cleared_input"))

    def _clear_output(self):
        self.converted_files = []
        self.result_text.configure(state="normal")
        self.result_text.delete("1.0", "end")
        self.result_text.insert("1.0", self._t("waiting_convert"))
        self.result_text.configure(state="disabled")
        self.progress.set(0)
        self.progress_label.configure(text=self._t("cleared_output"))

    def _browse_output(self):
        folder = filedialog.askdirectory(title=self._t("browse"))
        if folder:
            self.output_entry.delete(0, "end")
            self.output_entry.insert(0, folder)


    def _start_convert(self):
        if not self.source_files:
            messagebox.showwarning(self._t("hint"), self._t("please_add_files_convert"))
            return

        target = self.target_format.get().lower().strip()
        if not target:
            messagebox.showwarning(self._t("hint"), self._t("please_select_format"))
            return

        out_dir = self.output_entry.get().strip()
        if not out_dir:
            out_dir = self._get_default_output_dir()
        os.makedirs(out_dir, exist_ok=True)

        self.result_text.configure(state="normal")
        self.result_text.delete("1.0", "end")
        self.result_text.configure(state="disabled")

        self.progress.set(0)
        self.progress_label.configure(text=self._t("progress") + ": 0%")
        self.start_btn.configure(state="disabled")

        supported_files = []
        skipped = []
        for fp in self.source_files:
            if os.path.exists(fp):
                ext = _get_ext(fp)
                allowed = get_allowed_formats(ext)
                if target in allowed:
                    supported_files.append(fp)
                else:
                    skipped.append(fp)
            else:
                skipped.append(fp)

        result_buffer = []
        if skipped:
            result_buffer.append(self._t("skip_unsupported") + "\n")
            for fp in skipped:
                result_buffer.append(f"  - {os.path.basename(fp)}\n")
            result_buffer.append("\n")

        if not supported_files:
            result_buffer.append(self._t("please_add_files_convert") + "\n")
            self._append_result_batch(result_buffer)
            self.start_btn.configure(state="normal")
            return

        total = len(supported_files)
        result_buffer.append(f"\n{self._t('start_convert')}: {total} {self._t('files')}\n")
        result_buffer.append("=" * 50 + "\n")
        self._append_result_batch(result_buffer)
        result_buffer = []

        # 分离 COM 转换文件（COM 非线程安全，需单线程处理）
        com_exts = ("doc", "docx", "xls", "xlsx", "ppt", "pptx")
        com_files = [fp for fp in supported_files if _get_ext(fp) in com_exts and win32com_client]
        normal_files = [fp for fp in supported_files if fp not in com_files]

        result_queue = queue.Queue()
        success_count = [0]
        failed_count = [0]
        completed_count = [0]
        lock = threading.Lock()
        stop_event = threading.Event()

        def _process_one(fp):
            try:
                result = convert_file(fp, out_dir, target)
                with lock:
                    success_count[0] += 1
                result_queue.put({"status": "success", "file": fp, "result": result})
            except ConversionError as e:
                with lock:
                    failed_count[0] += 1
                result_queue.put({"status": "error", "file": fp, "error": str(e)})
            except Exception as e:
                import traceback
                with lock:
                    failed_count[0] += 1
                result_queue.put({"status": "error", "file": fp, "error": f"{e}\n{traceback.format_exc()}"})

        def _ui_updater():
            """批量刷新 UI，每 150ms 执行一次"""
            batch = []
            while not result_queue.empty():
                try:
                    batch.append(result_queue.get_nowait())
                except queue.Empty:
                    break
            if batch:
                texts = []
                for res in batch:
                    if res["status"] == "success":
                        texts.append(f"[{os.path.basename(res['file'])}] -> [{target}] OK\n")
                        texts.append(f"  {self._t('saved_to')}{res['result']}\n")
                        self.converted_files.append(res["result"])
                    else:
                        texts.append(f"[{os.path.basename(res['file'])}] FAILED\n")
                        texts.append(f"  {self._t('error')}: {res['error']}\n")
                self._append_result_batch(texts)
                completed_count[0] += len(batch)
                pct = completed_count[0] / total
                self.progress.set(pct)
                self.progress_label.configure(
                    text=f"{self._t('progress')}: {completed_count[0]}/{total} ({int(pct*100)}%)"
                )
            if not stop_event.is_set() or not result_queue.empty():
                self.after(150, _ui_updater)

        # 启动 UI 刷新定时器
        self.after(150, _ui_updater)

        def _convert_worker():
            # 先处理非 COM 文件（多线程）
            if normal_files:
                with ThreadPoolExecutor(max_workers=min(4, len(normal_files))) as executor:
                    list(executor.map(_process_one, normal_files))
            # 再处理 COM 文件（单线程，避免 COM 冲突）
            for fp in com_files:
                _process_one(fp)
            stop_event.set()
            # 等 UI 刷新完成后再通知结束
            self.after(300, lambda: self._finish_convert(success_count[0], failed_count[0], total, out_dir))

        # 后台线程执行转换，主线程保持空闲处理 UI 事件
        threading.Thread(target=_convert_worker, daemon=True).start()

    def _finish_convert(self, success_count, failed_count, total, out_dir):
        self._append_result_batch([
            "=" * 50 + "\n",
            f"{self._t('convert_done')}: {success_count} {self._t('files')}, {self._t('failed_count')}: {failed_count} {self._t('files')}\n"
        ])
        self.progress.set(1.0 if total > 0 else 0)
        self.progress_label.configure(text=f"{self._t('progress')}: {total}/{total} (100%)")
        if success_count > 0:
            self._show_success_dialog(success_count, out_dir)
        self.start_btn.configure(state="normal")

    def _append_result(self, text):
        self._append_result_batch([text])

    def _append_result_batch(self, texts):
        if not texts:
            return
        self.result_text.configure(state="normal")
        self.result_text.insert("end", "".join(texts))
        # 限制结果文本行数，防止内存无限增长导致卡顿
        try:
            line_count = int(self.result_text.index("end-1c").split(".")[0])
            if line_count > 2000:
                self.result_text.delete("1.0", f"{line_count - 1500}.0")
        except Exception:
            pass
        self.result_text.configure(state="disabled")
        self.result_text.see("end")

    def _show_success_dialog(self, count, out_dir):
        dialog = CTkToplevel(self)
        dialog.title(self._t("success"))
        dialog.geometry("400x200")
        dialog.transient(self)
        dialog.lift()
        dialog.focus_force()
        dialog.grab_set()

        CTkLabel(dialog, text=f"{self._t('convert_done')}: {count} {self._t('files')}", font=("Microsoft YaHei", 16, "bold")).pack(pady=20)

        btn_frame = CTkFrame(dialog, fg_color="transparent")
        btn_frame.pack(pady=20)

        CTkButton(btn_frame, text=self._t("open_dir"), command=lambda: os.startfile(out_dir) if os.path.exists(out_dir) else None).pack(side="left", padx=10)
        CTkButton(btn_frame, text=self._t("close"), command=dialog.destroy).pack(side="left", padx=10)

    def _show_format_help(self):
        dialog = CTkToplevel(self)
        dialog.title(self._t("supported_formats"))
        dialog.geometry("700x500")
        dialog.transient(self)
        dialog.lift()
        dialog.focus_force()

        text = CTkTextbox(dialog, wrap="word")
        text.pack(fill="both", expand=True, padx=10, pady=10)
        text.insert("1.0", "\n" + self._t("supported_formats") + "\n")
        text.insert("end", "=" * 60 + "\n\n")

        for fmt, desc in sorted(FORMAT_DESCRIPTIONS.items()):
            targets = CONVERSION_MAP.get(fmt, [])
            targets_str = ", ".join(targets) if targets else "无"
            text.insert("end", f"【{fmt.upper()}】 {desc}\n")
            text.insert("end", f"    可转换至: {targets_str}\n\n")
        text.configure(state="disabled")

        CTkButton(dialog, text=self._t("close"), command=dialog.destroy).pack(pady=10)

    def _paste_sample(self):
        sample_dir = os.path.join(tempfile.gettempdir(), "transfer_samples")
        os.makedirs(sample_dir, exist_ok=True)

        sample_files = []
        txt_path = os.path.join(sample_dir, "sample.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("这是示例文本文件\n用于演示文件转换功能\n\n支持多种格式互转！")
        sample_files.append(txt_path)

        if Document:
            docx_path = os.path.join(sample_dir, "sample.docx")
            doc = Document()
            doc.add_heading("示例文档", level=1)
            doc.add_paragraph("这是一个示例 Word 文档，用于演示转换功能。")
            doc.add_paragraph("支持转换为 PDF、TXT、HTML 等多种格式。")
            doc.save(docx_path)
            sample_files.append(docx_path)

        csv_path = os.path.join(sample_dir, "sample.csv")
        with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            writer.writerow(["姓名", "年龄", "城市"])
            writer.writerow(["张三", "25", "北京"])
            writer.writerow(["李四", "30", "上海"])
            writer.writerow(["王五", "28", "广州"])
        sample_files.append(csv_path)

        if Image:
            img_path = os.path.join(sample_dir, "sample.png")
            img = Image.new("RGB", (200, 100), color="lightblue")
            img.save(img_path)
            sample_files.append(img_path)

        self.source_files.extend(sample_files)
        self._update_files_list()


def main():
    app = FileTransferApp()
    app.mainloop()


if __name__ == "__main__":
    main()
